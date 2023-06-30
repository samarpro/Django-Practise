from pathlib import Path
from django.shortcuts import render
from .models import TodoLists,FileUpload
from django.http import HttpResponseRedirect,HttpResponse
from .forms import CreateForm,DelForm,UploadFile,AccessForm
from docx import Document
import os.path

# varaible for base directory
# Create your views here.
def idirect(req,id):
    #object of TodoLists
    t=TodoLists.objects.get(id=id)
    return HttpResponse(f"{t} {t.id} {t.name}")
def home(req):
    Todos = TodoLists.objects.all()
    context={
        "Todos":Todos,
    }
    return render(req,"main/home.html",context)

def create(req):
    info = req.POST 
    if req.method =="POST":
        form = CreateForm(req.POST)
        if form.is_valid():
            n = form.cleaned_data['name']
            todo = TodoLists(name=n)
            todo.save()
            return HttpResponseRedirect(f"/{todo.id}")
    else:
        info = req.POST 

        form = CreateForm()
    context={'form':form,"info":info}
    return render(req,"main/create.html",context)

def delete(req):
    form = DelForm(req.POST)
    if form.is_valid():
        id  = form.cleaned_data['id_val']
        t = TodoLists.objects.get(id=id)
        t.delete() 
        return HttpResponseRedirect("/")
    context = {
        "form":form, 
        # "info":info
        }
    return render(req,"main/del.html",context)

def upload(req):
    if req.method == "POST":
        BASE_DIR = Path(__file__).resolve().parent.parent
        form = UploadFile(req.POST,req.FILES)
        if form.is_valid():
            text = form.cleaned_data['text']
            file = form.cleaned_data['file_path']
            f = FileUpload(text=text,fil=file)
            f.save()  
            wordFile = Document(os.path.join(BASE_DIR,f'media\word\{file.name}'))
            return HttpResponse(len(wordFile.paragraphs)) 
    else: 
        form = UploadFile() 
    context ={"form":form}
    return render(req,"main/upload.html",context)   

def access(req):
    if req.method=='POST':
        form = AccessForm(req.POST)
        if form.is_valid():
            BASE_DIR = Path(__file__).resolve().parent.parent
            text = form.cleaned_data['NameText']
            wordFile = Document(os.path.join(BASE_DIR,f'media\word\{text}'))
            return HttpResponse(len(wordFile.paragraphs)) 

    else:
        form = AccessForm()
    context = {"form":form}

    return render(req,"main/access.html",context)

    